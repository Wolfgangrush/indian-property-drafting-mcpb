"""Build a transactional-instrument reference.docx for the contracts + property
drafting plugins.

These plugins produce contracts / conveyancing instruments, NOT pleadings.
The styling convention differs from pleadings:

- Paper: A4
- Body (Normal): TNR 12pt, single line spacing, justified, no first-line indent
- Heading 1 (Document title / Major Recital): TNR 14pt bold centered, NOT underlined
- Heading 2 (Section title — e.g., "1. DEFINITIONS"): TNR 12pt bold left, no underline
- Heading 3 (Sub-section): TNR 12pt bold + italic left
- Heading 4 (Sub-sub-section): TNR 12pt italic left
- Margins: 2.5cm all around (narrower binding side than pleadings)
- Tables: tblLayout=fixed; first row bold; cell margins locked
- No letter-spacing on Heading 2 (no spaced "F A C T S" effect — wrong register for contracts)
- No underlining of headings (commercial document convention)
- No "MOST RESPECTFULLY SHEWETH:" anchor
"""
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("reference.docx")


def lock_style(style, *, font="Times New Roman", size_pt=12, bold=False,
               italic=False, underline=False, align=None, line_spacing=1.0,
               line_spacing_rule=WD_LINE_SPACING.SINGLE,
               color="000000", letter_spacing_pt=None,
               space_before_pt=0, space_after_pt=6,
               first_line_indent_cm=None, keep_with_next=False,
               outline_level=None):
    font_obj = style.font
    font_obj.name = font
    font_obj.size = Pt(size_pt)
    font_obj.bold = bold
    font_obj.italic = italic
    if underline:
        font_obj.underline = True
    if color:
        font_obj.color.rgb = RGBColor.from_string(color)

    rpr = style.element.get_or_add_rPr()
    existing = rpr.find(qn("w:rFonts"))
    if existing is not None:
        rpr.remove(existing)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rpr.append(rFonts)

    if letter_spacing_pt is not None:
        spacing_el = OxmlElement("w:spacing")
        spacing_el.set(qn("w:val"), str(int(letter_spacing_pt * 20)))
        rpr.append(spacing_el)

    if underline:
        u_el = OxmlElement("w:u")
        u_el.set(qn("w:val"), "single")
        rpr.append(u_el)

    pf = style.paragraph_format
    if align is not None:
        pf.alignment = align
    pf.line_spacing = line_spacing
    pf.line_spacing_rule = line_spacing_rule
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    pf.keep_with_next = keep_with_next

    if outline_level is not None:
        ppr = style.element.get_or_add_pPr()
        existing_ol = ppr.find(qn("w:outlineLvl"))
        if existing_ol is not None:
            ppr.remove(existing_ol)
        ol = OxmlElement("w:outlineLvl")
        ol.set(qn("w:val"), str(outline_level))
        ppr.append(ol)


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Normal — TNR 12pt single-spaced justified
    normal = doc.styles["Normal"]
    lock_style(normal, font="Times New Roman", size_pt=12, bold=False,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.0,
               line_spacing_rule=WD_LINE_SPACING.SINGLE,
               first_line_indent_cm=0, space_after_pt=6)

    # Heading 1 — TNR 14pt bold centered (document title)
    h1 = doc.styles["Heading 1"]
    lock_style(h1, font="Times New Roman", size_pt=14, bold=True,
               underline=False, align=WD_ALIGN_PARAGRAPH.CENTER,
               line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE,
               space_before_pt=12, space_after_pt=12, keep_with_next=True,
               outline_level=0)
    h1.paragraph_format.first_line_indent = Cm(0)

    # Heading 2 — TNR 12pt bold left (section title, e.g. "1. DEFINITIONS")
    h2 = doc.styles["Heading 2"]
    lock_style(h2, font="Times New Roman", size_pt=12, bold=True,
               underline=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE,
               space_before_pt=12, space_after_pt=6, keep_with_next=True,
               outline_level=1)
    h2.paragraph_format.first_line_indent = Cm(0)

    # Heading 3 — TNR 12pt bold italic left (sub-section)
    h3 = doc.styles["Heading 3"]
    lock_style(h3, font="Times New Roman", size_pt=12, bold=True,
               italic=True, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE,
               space_before_pt=6, space_after_pt=6, keep_with_next=True,
               outline_level=2)
    h3.paragraph_format.first_line_indent = Cm(0)

    # Heading 4 — TNR 12pt italic left (sub-sub-section)
    h4 = doc.styles["Heading 4"]
    lock_style(h4, font="Times New Roman", size_pt=12, bold=False,
               italic=True, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT,
               line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE,
               space_before_pt=6, space_after_pt=6, keep_with_next=True,
               outline_level=3)
    h4.paragraph_format.first_line_indent = Cm(0)

    # Title style — for deed / contract title
    title = doc.styles["Title"]
    lock_style(title, font="Times New Roman", size_pt=14, bold=True,
               underline=False, align=WD_ALIGN_PARAGRAPH.CENTER,
               line_spacing=1.0, line_spacing_rule=WD_LINE_SPACING.SINGLE,
               space_before_pt=0, space_after_pt=12, keep_with_next=True)
    title.paragraph_format.first_line_indent = Cm(0)

    # Page numbers at bottom center (commercial document convention)
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    # Demo content
    doc.add_paragraph("THIS DEED OF [ASSIGNMENT/SALE/LEASE/etc.]", style="Heading 1")
    doc.add_paragraph("(Transactional instrument — TNR 12pt single-spaced; "
                      "this is a style template, pandoc replaces this content "
                      "when rendering the actual document.)")
    doc.add_paragraph("1. PARTIES", style="Heading 2")
    doc.add_paragraph("This Deed is executed on this ___ day of _____, 20___ "
                      "by and between …")
    doc.add_paragraph("2. DEFINITIONS", style="Heading 2")
    doc.add_paragraph("2.1 Definitions", style="Heading 3")
    doc.add_paragraph("In this Deed, unless the context otherwise requires, the "
                      "following terms shall have the meanings ascribed to them …")

    doc.save(str(OUT_PATH))
    print(f"OK · wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
